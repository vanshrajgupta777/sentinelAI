"""
Document Processing
===================
Real local text extraction. Supports:
  * PDF (text PDFs via pypdf; scanned PDFs flagged for OCR)
  * DOCX
  * Plain text files
  * Code archives (.zip) — enumerated entry list, no execution
  * Image files — basic metadata + OCR-not-implemented flag

Storage is local (workspaces/). No external calls.
"""
from __future__ import annotations

import io
import logging
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

logger = logging.getLogger("sentinelai.docs")


@dataclass
class ExtractedDocument:
    text: str
    page_count: int
    ocr_required: bool
    metadata: dict


def _safe_extract_pdf(path: Path) -> ExtractedDocument:
    text_parts: List[str] = []
    pages = 0
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = len(reader.pages)
        for i, page in enumerate(reader.pages):
            try:
                t = page.extract_text() or ""
            except Exception as e:  # pragma: no cover
                t = ""
                logger.warning("PDF page %d extraction failed: %s", i, e)
            if t:
                text_parts.append(f"[Page {i + 1}]\n{t}")
        text = "\n\n".join(text_parts)
        # Heuristic: if the extracted text is empty or very short and
        # the file is a PDF, flag for OCR.
        ocr_required = len(text.strip()) < 20
    except Exception as e:
        logger.warning("PDF parsing failed: %s", e)
        text = ""
        pages = 0
        ocr_required = True
    return ExtractedDocument(text=text, page_count=pages, ocr_required=ocr_required, metadata={"format": "pdf"})


def _safe_extract_docx(path: Path) -> ExtractedDocument:
    try:
        from docx import Document as DocxDocument

        d = DocxDocument(str(path))
        paragraphs = [p.text for p in d.paragraphs if p.text]
        text = "\n".join(paragraphs)
        return ExtractedDocument(
            text=text,
            page_count=max(1, len(text) // 2000),
            ocr_required=False,
            metadata={"format": "docx", "paragraphs": len(paragraphs)},
        )
    except Exception as e:
        logger.warning("DOCX parsing failed: %s", e)
        return ExtractedDocument(text="", page_count=0, ocr_required=False, metadata={"format": "docx", "error": str(e)})


def _safe_extract_text(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(errors="ignore")
        return ExtractedDocument(
            text=text,
            page_count=max(1, len(text) // 2000),
            ocr_required=False,
            metadata={"format": "text", "chars": len(text)},
        )
    except Exception as e:
        return ExtractedDocument(text="", page_count=0, ocr_required=False, metadata={"format": "text", "error": str(e)})


def _safe_extract_zip(path: Path) -> ExtractedDocument:
    """List entries in a code archive. Never executes anything."""
    try:
        with zipfile.ZipFile(str(path)) as z:
            names = z.namelist()
            sample_lines: List[str] = []
            for n in names[:25]:
                if n.endswith(("/", "\\")):
                    continue
                try:
                    with z.open(n) as f:
                        head = f.read(4096).decode("utf-8", errors="ignore")
                    if head.strip():
                        sample_lines.append(f"--- {n} ---")
                        sample_lines.append(head[:1500])
                except Exception:
                    continue
            text = "\n".join(sample_lines)
            return ExtractedDocument(
                text=text or "(empty archive)",
                page_count=1,
                ocr_required=False,
                metadata={"format": "zip", "entries": len(names)},
            )
    except Exception as e:
        return ExtractedDocument(text="", page_count=0, ocr_required=False, metadata={"format": "zip", "error": str(e)})


def _safe_extract_image(path: Path) -> ExtractedDocument:
    return ExtractedDocument(
        text="",
        page_count=1,
        ocr_required=True,
        metadata={"format": "image", "ocr_pending": True},
    )


def extract(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _safe_extract_pdf(path)
    if suffix == ".docx":
        return _safe_extract_docx(path)
    if suffix in {".txt", ".md", ".csv", ".log"}:
        return _safe_extract_text(path)
    if suffix in {".zip"}:
        return _safe_extract_zip(path)
    if suffix in {".png", ".jpg", ".jpeg", ".dwg"}:
        return _safe_extract_image(path)
    # Unknown — best effort as text
    return _safe_extract_text(path)


def save_upload(upload_bytes: bytes, dest_dir: Path, name: str) -> Path:
    dest_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", name)
    out = dest_dir / safe_name
    with open(out, "wb") as f:
        f.write(upload_bytes)
    return out
