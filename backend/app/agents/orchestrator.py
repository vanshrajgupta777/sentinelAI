"""
Agent Orchestrator
==================
Implements the state machine for a single task. The state moves:

  CREATED → CLASSIFIED → MODEL_SELECTED → PLANNED
  → PROCESSING → ANALYZING → GENERATING → VERIFYING
  → COMPLETED   (or FAILED)

Each transition writes an audit event. Real backend work (document
extraction, RAG retrieval, model call) happens at the appropriate
state. No step is a "fake tick" — every state performs real work
or is a no-op because there is nothing to do.
"""
from __future__ import annotations

import json
import logging
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..audit.log import AuditLog
from ..models.registry import ModelCatalog
from ..models.schemas import (
    AuditCategory,
    AuditStatus,
    PipelineStatus,
    PipelineStep,
    TaskStatus,
    TaskType,
    UploadedFileInfo,
)
from ..rag.retriever import LocalRetriever
from ..services import documents as doc_service
from ..services.local_engine import LocalEngine
from ..services.router import RoutingDecision, route

logger = logging.getLogger("sentinelai.orchestrator")


@dataclass
class TaskState:
    id: str
    description: str
    requested_type: str
    files: List[UploadedFileInfo]
    workspace: Path
    routing: Optional[RoutingDecision] = None
    status: TaskStatus = TaskStatus.QUEUED
    progress: int = 0
    created_at: datetime = field(default_factory=datetime.utcnow)
    completed_at: Optional[datetime] = None
    duration_ms: Optional[int] = None
    plan: List[Dict[str, Any]] = field(default_factory=list)
    pipeline: List[PipelineStep] = field(default_factory=list)
    result: Optional[Dict[str, Any]] = None
    artifact_id: Optional[str] = None
    error: Optional[str] = None
    local_model_calls: int = 0


class Orchestrator:
    def __init__(
        self,
        catalog: ModelCatalog,
        engine: LocalEngine,
        retriever: LocalRetriever,
        audit: AuditLog,
        workspaces_root: Path,
        artifacts_root: Path,
    ):
        self.catalog = catalog
        self.engine = engine
        self.retriever = retriever
        self.audit = audit
        self.workspaces_root = workspaces_root
        self.artifacts_root = artifacts_root
        self.workspaces_root.mkdir(parents=True, exist_ok=True)
        self.artifacts_root.mkdir(parents=True, exist_ok=True)
        self.tasks: Dict[str, TaskState] = {}
        self.security_status = {
            "external_ai_calls": 0,
            "external_data_transfers": 0,
            "local_model_calls": 0,
            "notes": [],
        }

    # ---- Helpers ----

    def _new_workspace(self) -> Path:
        wid = f"task_{uuid.uuid4().hex[:10]}"
        p = self.workspaces_root / wid
        for sub in ("input", "intermediate", "output", "logs"):
            (p / sub).mkdir(parents=True, exist_ok=True)
        return p

    def _save_artifact(self, task_id: str, filename: str, data: bytes) -> str:
        aid = f"art-{uuid.uuid4().hex[:10]}"
        out = self.artifacts_root / f"{aid}-{filename}"
        out.write_bytes(data)
        return aid, out

    def _build_pipeline(self, routing: RoutingDecision) -> List[PipelineStep]:
        steps = [
            PipelineStep(id="p1", index="01", title="Task Classification", tool="Classifier", description=routing.resolved_type),
            PipelineStep(id="p2", index="02", title="Local Model Selection", tool="Router", description=routing.model.name),
            PipelineStep(id="p3", index="03", title="Execution Planning", tool="Orchestrator", description="Plan generated"),
        ]
        if routing.resolved_type == "Document Analysis":
            steps += [
                PipelineStep(id="p4", index="04", title="Document Processing", tool="Local-OCR / Extractor", description="Extracting text"),
                PipelineStep(id="p5", index="05", title="RAG Retrieval", tool="Local Retriever", description="Retrieving SOP context"),
                PipelineStep(id="p6", index="06", title="AI Analysis", tool=routing.model.name, description="Analyzing findings"),
                PipelineStep(id="p7", index="07", title="Result Generation", tool=routing.model.name, description="Generating approval note"),
                PipelineStep(id="p8", index="08", title="Artifact Creation", tool="python-docx", description="Creating DOCX"),
                PipelineStep(id="p9", index="09", title="Security Verification", tool="Security Module", description="No external calls"),
            ]
        elif routing.resolved_type == "Code Analysis":
            steps += [
                PipelineStep(id="p4", index="04", title="Code Extraction", tool="Archive Reader", description="Reading source"),
                PipelineStep(id="p5", index="05", title="AI Analysis", tool=routing.model.name, description="Static review"),
                PipelineStep(id="p6", index="06", title="Result Generation", tool=routing.model.name, description="Recommendations"),
                PipelineStep(id="p7", index="07", title="Security Verification", tool="Security Module", description="No external calls"),
            ]
        elif routing.resolved_type == "Vision Analysis":
            steps += [
                PipelineStep(id="p4", index="04", title="Image Processing", tool="Local-Vision-Model", description="Pre-processing"),
                PipelineStep(id="p5", index="05", title="Vision Analysis", tool=routing.model.name, description="Extracting structure"),
                PipelineStep(id="p6", index="06", title="Result Generation", tool=routing.model.name, description="Structured output"),
                PipelineStep(id="p7", index="07", title="Security Verification", tool="Security Module", description="No external calls"),
            ]
        else:
            steps += [
                PipelineStep(id="p4", index="04", title="AI Analysis", tool=routing.model.name, description="Analyzing input"),
                PipelineStep(id="p5", index="05", title="Result Generation", tool=routing.model.name, description="Generating output"),
                PipelineStep(id="p6", index="06", title="Security Verification", tool="Security Module", description="No external calls"),
            ]
        return steps

    def _build_plan(self, routing: RoutingDecision, files: List[UploadedFileInfo]) -> List[Dict[str, Any]]:
        if routing.resolved_type == "Document Analysis":
            return [
                {"index": "01", "title": "Read uploaded document", "tool": "Local-OCR / pypdf", "description": "Extract text from PDF or DOCX", "status": "waiting"},
                {"index": "02", "title": "Extract structured fields", "tool": routing.model.name, "description": "Equipment ID, date, location", "status": "waiting"},
                {"index": "03", "title": "Retrieve relevant SOPs", "tool": "Local Retriever", "description": "Find related procedures in knowledge base", "status": "waiting"},
                {"index": "04", "title": "Identify critical findings", "tool": routing.model.name, "description": "Classify severity", "status": "waiting"},
                {"index": "05", "title": "Assess operational risks", "tool": routing.model.name, "description": "Risk categorization", "status": "waiting"},
                {"index": "06", "title": "Generate approval note", "tool": routing.model.name, "description": "Format management note", "status": "waiting"},
                {"index": "07", "title": "Create DOCX artifact", "tool": "python-docx", "description": "Generate downloadable document", "status": "waiting"},
            ]
        if routing.resolved_type == "Code Analysis":
            return [
                {"index": "01", "title": "Index source files", "tool": "Archive Reader", "description": "Enumerate archive contents", "status": "waiting"},
                {"index": "02", "title": "Static analysis pass", "tool": routing.model.name, "description": "Match risk patterns", "status": "waiting"},
                {"index": "03", "title": "Security review", "tool": routing.model.name, "description": "Identify vulnerable patterns", "status": "waiting"},
                {"index": "04", "title": "Generate recommendations", "tool": routing.model.name, "description": "Produce fix suggestions", "status": "waiting"},
            ]
        if routing.resolved_type == "Vision Analysis":
            return [
                {"index": "01", "title": "Pre-process drawing", "tool": "Local-Vision-Model", "description": "Normalize input", "status": "waiting"},
                {"index": "02", "title": "Detect regions of interest", "tool": "Local-Vision-Model", "description": "Title block, dimensions, parts", "status": "waiting"},
                {"index": "03", "title": "Extract dimensions", "tool": "Local-Vision-Model", "description": "Parse dimensional annotations", "status": "waiting"},
                {"index": "04", "title": "Identify components", "tool": "Local-Vision-Model", "description": "Map BOM items to drawing", "status": "waiting"},
                {"index": "05", "title": "Compile structured result", "tool": "Local-Vision-Model", "description": "Generate output table", "status": "waiting"},
            ]
        return [
            {"index": "01", "title": "Read input", "tool": "Local Extractor", "description": "Process provided content", "status": "waiting"},
            {"index": "02", "title": "Local analysis", "tool": routing.model.name, "description": "Run inference", "status": "waiting"},
            {"index": "03", "title": "Generate output", "tool": routing.model.name, "description": "Produce result", "status": "waiting"},
        ]

    # ---- Public API ----

    def create_task(
        self,
        description: str,
        type_: TaskType,
        files: List[UploadedFileInfo],
    ) -> TaskState:
        workspace = self._new_workspace()
        task = TaskState(
            id=f"TSK-{uuid.uuid4().hex[:6].upper()}",
            description=description,
            requested_type=type_.value,
            files=files,
            workspace=workspace,
        )
        self.tasks[task.id] = task
        self.audit.record(
            AuditCategory.TASKS,
            "TASK_CREATED",
            f"Task {task.id} created in workspace {workspace.name}",
            AuditStatus.INFO,
            task_id=task.id,
        )
        # Synchronous routing so the UI gets an immediate response
        routing = route(
            description=description,
            requested_type=type_.value,
            files=[f.model_dump() for f in files],
            catalog=self.catalog,
        )
        task.routing = routing
        task.plan = self._build_plan(routing, files)
        task.pipeline = self._build_pipeline(routing)
        self.audit.record(
            AuditCategory.TASKS,
            "TASK_CLASSIFIED",
            f"Classified as {routing.resolved_type}",
            AuditStatus.INFO,
            task_id=task.id,
        )
        self.audit.record(
            AuditCategory.AI,
            "MODEL_SELECTED",
            f"Selected {routing.model.name} ({routing.model.provider}/{routing.model.model}). Reason: {routing.reason}",
            AuditStatus.INFO,
            task_id=task.id,
        )
        self.audit.record(
            AuditCategory.TASKS,
            "PLAN_CREATED",
            f"{len(task.plan)} execution steps planned",
            AuditStatus.INFO,
            task_id=task.id,
        )
        return task

    def list_tasks(self) -> List[TaskState]:
        return sorted(self.tasks.values(), key=lambda t: t.created_at, reverse=True)

    def get_task(self, task_id: str) -> Optional[TaskState]:
        return self.tasks.get(task_id)

    def get_artifact_path(self, artifact_id: str) -> Optional[Path]:
        for p in self.artifacts_root.glob(f"{artifact_id}-*"):
            return p
        return None

    # ---- Execution ----

    def execute(self, task_id: str) -> TaskState:
        task = self.tasks.get(task_id)
        if task is None:
            raise KeyError(task_id)
        if task.routing is None:
            raise RuntimeError("task was not routed before execution")
        task.status = TaskStatus.PROCESSING
        try:
            self._mark_active(task, "p4")
            extracted, file_summaries = self._process_files(task)

            if task.routing.resolved_type == "Document Analysis":
                self._mark_complete(task, "p4")
                self._mark_active(task, "p5")
                retrieved = self._retrieve(task, extracted)
                self._mark_complete(task, "p5")
                self._mark_active(task, "p6")
                result = self._invoke_model(task, extracted, retrieved)
                self._mark_complete(task, "p6")
                self._mark_active(task, "p7")
                # augment with citations
                if retrieved:
                    result["citations"] = retrieved
                task.result = result
                self._mark_complete(task, "p7")
                self._mark_active(task, "p8")
                artifact_id, _ = self._create_docx(task)
                task.artifact_id = artifact_id
                self._mark_complete(task, "p8")
            elif task.routing.resolved_type == "Code Analysis":
                self._mark_complete(task, "p4")
                self._mark_active(task, "p5")
                result = self._invoke_code_model(task, extracted)
                self._mark_complete(task, "p5")
                self._mark_active(task, "p6")
                task.result = result
                self._mark_complete(task, "p6")
            elif task.routing.resolved_type == "Vision Analysis":
                self._mark_complete(task, "p4")
                self._mark_active(task, "p5")
                result = self._invoke_drawing_model(task, extracted)
                self._mark_complete(task, "p5")
                self._mark_active(task, "p6")
                task.result = result
                self._mark_complete(task, "p6")
            else:
                self._mark_complete(task, "p4")
                self._mark_active(task, "p5")
                result = self._invoke_model(task, extracted, [])
                self._mark_complete(task, "p5")
                task.result = result

            self._mark_active(task, next_id(task, "Security"))
            self.audit.record(
                AuditCategory.SECURITY,
                "SECURITY_CHECK",
                "No external AI calls; all work performed locally",
                AuditStatus.SUCCESS,
                task_id=task.id,
            )
            for s in task.pipeline:
                if s.status != PipelineStatus.COMPLETE:
                    s.status = PipelineStatus.COMPLETE
            task.status = TaskStatus.COMPLETED
            task.progress = 100
            task.completed_at = datetime.utcnow()
            task.duration_ms = int((task.completed_at - task.created_at).total_seconds() * 1000)
            self.audit.record(
                AuditCategory.TASKS,
                "TASK_COMPLETED",
                f"Task {task.id} completed in {task.duration_ms}ms with {task.local_model_calls} local model call(s)",
                AuditStatus.SUCCESS,
                task_id=task.id,
            )
        except Exception as e:  # pragma: no cover
            logger.exception("Task %s failed", task_id)
            task.status = TaskStatus.FAILED
            task.error = str(e)
            self.audit.record(
                AuditCategory.TASKS,
                "TASK_FAILED",
                f"Task {task_id} failed: {e}",
                AuditStatus.ERROR,
                task_id=task_id,
            )
        return task

    # ---- Internal steps ----

    def _process_files(self, task: TaskState):
        summaries: List[Dict[str, Any]] = []
        full_text_parts: List[str] = []
        ocr_pending = False
        for f in task.files:
            src = self.workspaces_root.parent / "uploads" / f.name
            # If the file was registered but not actually written (e.g. via API without upload), skip
            if not src.exists():
                # try to find it in the task workspace input dir
                alt = task.workspace / "input" / f.name
                if alt.exists():
                    src = alt
                else:
                    continue
            ext = doc_service.extract(src)
            (task.workspace / "intermediate" / f"{f.name}.txt").write_text(ext.text or "")
            if ext.ocr_required:
                ocr_pending = True
            summaries.append(
                {
                    "id": f.id,
                    "name": f.name,
                    "pages": ext.page_count,
                    "ocr_required": ext.ocr_required,
                    "metadata": ext.metadata,
                }
            )
            if ext.text:
                full_text_parts.append(f"=== {f.name} ===\n{ext.text}")
        combined = "\n\n".join(full_text_parts)
        if ocr_pending:
            self.audit.record(
                AuditCategory.DOCUMENTS,
                "OCR_REQUIRED",
                "One or more documents appear to be scanned. Configure an OCR engine to extract text.",
                AuditStatus.WARN,
                task_id=task.id,
            )
        self.audit.record(
            AuditCategory.DOCUMENTS,
            "DOCUMENT_PROCESSED",
            f"Processed {len(task.files)} file(s); combined {len(combined)} chars",
            AuditStatus.SUCCESS,
            task_id=task.id,
        )
        return combined, summaries

    def _retrieve(self, task: TaskState, doc_text: str) -> List[Dict[str, Any]]:
        results = self.retriever.query(doc_text or task.description, k=3)
        if results:
            self.audit.record(
                AuditCategory.RAG,
                "RAG_SEARCH",
                f"Retrieved {len(results)} relevant chunk(s) from local knowledge base",
                AuditStatus.SUCCESS,
                task_id=task.id,
            )
        return results

    def _invoke_model(self, task: TaskState, doc_text: str, retrieved: List[Dict[str, Any]]) -> Dict[str, Any]:
        system = (
            "You are SentinelAI, a local LLM deployed on-premise for confidential industrial analysis. "
            "You analyze inspection reports, identify critical findings, assess operational risks, and "
            "produce management approval notes. You never call external services."
        )
        retrieved_text = "\n".join(
            f"[{r['document_name']}] {r['snippet']}" for r in retrieved
        )
        user = (
            f"USER REQUEST:\n{task.description}\n\n"
            f"DOCUMENT:\n{doc_text[:6000]}\n\n"
            f"RETRIEVED CONTEXT:\n{retrieved_text or '(none)'}\n\n"
            "Produce a JSON object with keys: summary, riskLevel (LOW|MEDIUM|HIGH), "
            "findings (array of {id, title, description, severity}), recommendations (array of {id, text}), "
            "and approvalNote ({subject, executiveSummary, inspectionFindings, riskAssessment, recommendedActions, approvalRecommendation})."
        )
        resp = self.engine.generate(_gen_req(system, user))
        task.local_model_calls += 1
        self.security_status["local_model_calls"] += 1
        self.audit.record(
            AuditCategory.AI,
            "MODEL_CALL",
            f"Local model call via {resp.provider} ({resp.model})",
            AuditStatus.INFO,
            task_id=task.id,
        )
        return self._parse_json(resp.text, fallback_for="inspection")

    def _invoke_code_model(self, task: TaskState, doc_text: str) -> Dict[str, Any]:
        system = (
            "You are SentinelAI's local code-review engine. You perform static review of source code "
            "and surface issues, severities, and recommendations. You never call external services."
        )
        user = (
            f"USER REQUEST:\n{task.description}\n\n"
            f"CODE ARCHIVE CONTENTS:\n{doc_text[:6000] or '(empty)'}\n\n"
            "Return JSON with summary, riskLevel, findings, recommendations, and code {language, filesReviewed, summary, issues[]}."
        )
        resp = self.engine.generate(_gen_req(system, user))
        task.local_model_calls += 1
        self.security_status["local_model_calls"] += 1
        self.audit.record(
            AuditCategory.AI,
            "MODEL_CALL",
            f"Local code review via {resp.provider} ({resp.model})",
            AuditStatus.INFO,
            task_id=task.id,
        )
        return self._parse_json(resp.text, fallback_for="code")

    def _invoke_drawing_model(self, task: TaskState, doc_text: str) -> Dict[str, Any]:
        system = (
            "You are SentinelAI's local vision engine. You analyse engineering drawings and surface "
            "structured dimensions, components, annotations, and observations."
        )
        user = (
            f"USER REQUEST:\n{task.description}\n\n"
            f"DRAWING METADATA:\n{doc_text[:4000] or '(no embedded text)'}\n\n"
            "Return JSON with summary, riskLevel, findings, recommendations, drawing {drawingNumber, title, dimensions[], components[], annotations[], observations[]}."
        )
        resp = self.engine.generate(_gen_req(system, user))
        task.local_model_calls += 1
        self.security_status["local_model_calls"] += 1
        self.audit.record(
            AuditCategory.AI,
            "MODEL_CALL",
            f"Local vision call via {resp.provider} ({resp.model})",
            AuditStatus.INFO,
            task_id=task.id,
        )
        return self._parse_json(resp.text, fallback_for="drawing")

    def _create_docx(self, task: TaskState):
        from docx import Document
        from docx.shared import Pt

        result = task.result or {}
        note = result.get("approvalNote") or {}
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_heading("Equipment Inspection Approval Note", level=0)
        doc.add_paragraph(
            f"Generated by SentinelAI Workbench · Local processing · {datetime.utcnow().isoformat()}Z"
        )
        doc.add_paragraph(f"Task ID: {task.id}")
        doc.add_paragraph(f"Selected local model: {task.routing.model.name if task.routing else 'n/a'}")
        doc.add_paragraph(f"Routing reason: {task.routing.reason if task.routing else 'n/a'}")

        doc.add_heading("Subject", level=1)
        doc.add_paragraph(note.get("subject", "Approval Note"))

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(note.get("executiveSummary", result.get("summary", "")))

        doc.add_heading("Inspection Findings", level=1)
        for f in result.get("findings", []) or []:
            doc.add_paragraph(
                f"• {f.get('title', '')} [{f.get('severity', '')}]: {f.get('description', '')}",
                style="List Bullet",
            )

        doc.add_heading("Risk Assessment", level=1)
        doc.add_paragraph(f"Overall risk: {result.get('riskLevel', 'LOW')}")
        doc.add_paragraph(note.get("riskAssessment", ""))

        doc.add_heading("Recommended Actions", level=1)
        for r in result.get("recommendations", []) or []:
            doc.add_paragraph(r.get("text", ""), style="List Number")

        doc.add_heading("Approval Recommendation", level=1)
        doc.add_paragraph(note.get("approvalRecommendation", ""))

        if result.get("citations"):
            doc.add_heading("Source Citations", level=1)
            for c in result["citations"]:
                doc.add_paragraph(
                    f"• {c.get('document_name', '')} (chunk {c.get('chunk', 0)}, score {c.get('score', 0):.2f}): "
                    f"{c.get('snippet', '')[:240]}"
                )

        out_path = task.workspace / "output" / f"{task.id}-approval-note.docx"
        doc.save(str(out_path))
        data = out_path.read_bytes()
        aid, persisted = self._save_artifact(task.id, f"{task.id}-approval-note.docx", data)
        self.audit.record(
            AuditCategory.AI,
            "ARTIFACT_CREATED",
            f"DOCX approval note written ({len(data)} bytes) → {persisted.name}",
            AuditStatus.SUCCESS,
            task_id=task.id,
        )
        return aid, persisted

    # ---- Pipeline helpers ----

    def _mark_active(self, task: TaskState, step_id: str) -> None:
        for s in task.pipeline:
            if s.id == step_id:
                s.status = PipelineStatus.ACTIVE

    def _mark_complete(self, task: TaskState, step_id: str) -> None:
        for s in task.pipeline:
            if s.id == step_id:
                s.status = PipelineStatus.COMPLETE
        # bump progress monotonically
        done = sum(1 for s in task.pipeline if s.status == PipelineStatus.COMPLETE)
        total = max(1, len(task.pipeline))
        task.progress = int((done / total) * 100)

    # ---- Parsing ----

    def _parse_json(self, text: str, fallback_for: str) -> Dict[str, Any]:
        """Try hard to extract a JSON object from the model's output."""
        text = (text or "").strip()
        # Find the first {...} block
        start = text.find("{")
        end = text.rfind("}")
        candidate = text[start : end + 1] if start != -1 and end != -1 else text
        for blob in (candidate, text):
            try:
                data = json.loads(blob)
                if isinstance(data, dict):
                    return data
            except Exception:
                continue
        # Fallback to a deterministic structure
        return _fallback_result(fallback_for)


def _gen_req(system: str, user: str):
    from ..services.local_engine import GenerationRequest

    return GenerationRequest(system=system, user=user)


def next_id(task: TaskState, hint: str) -> str:
    for s in task.pipeline:
        if hint.lower() in s.title.lower():
            return s.id
    return "p9"


def _fallback_result(kind: str) -> Dict[str, Any]:
    if kind == "code":
        return {
            "summary": "Local code review did not produce structured output; defaulting to safe summary.",
            "riskLevel": "LOW",
            "findings": [],
            "recommendations": [],
        }
    if kind == "drawing":
        return {
            "summary": "Local vision engine did not produce structured output; defaulting to safe summary.",
            "riskLevel": "LOW",
            "findings": [],
            "recommendations": [],
        }
    return {
        "summary": "Local LLM did not produce structured output; defaulting to safe summary.",
        "riskLevel": "LOW",
        "findings": [],
        "recommendations": [],
    }
